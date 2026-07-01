"""Document upload + extraction integration tests.

Exercises the full flow: synthesize a tiny PDF with reportlab, hit the
upload endpoint, kick the extract endpoint, poll until EXTRACTED, read
the extraction result.

The tests run with ``EVALS_MOCK=true`` so the agents return canned data
instead of calling Claude — keeps CI cheap and deterministic.
"""

from __future__ import annotations

import asyncio
import io
import os
import tempfile
from pathlib import Path
from uuid import UUID

import pytest

# Force SQLite + a temp DB before app imports — keeps tests isolated
# from any local fondok.db left behind by manual runs.
_TMP_DB = Path(tempfile.gettempdir()) / "fondok-tests-documents.db"
if _TMP_DB.exists():
    _TMP_DB.unlink()
os.environ["DATABASE_URL"] = f"sqlite+aiosqlite:///{_TMP_DB}"
os.environ["DOCUMENT_STORAGE_ROOT"] = str(
    Path(tempfile.gettempdir()) / "fondok-tests-storage"
)
os.environ.setdefault("EVALS_MOCK", "true")

# Wipe any leftover storage tree.
import shutil  # noqa: E402

_STORAGE_ROOT = Path(os.environ["DOCUMENT_STORAGE_ROOT"])
if _STORAGE_ROOT.exists():
    shutil.rmtree(_STORAGE_ROOT)


# ─────────────────────────── fixtures ───────────────────────────


def _build_sample_pdf() -> bytes:
    """Generate a 1-page hotel-themed PDF with reportlab.

    Stored in-memory so we never commit a binary blob. The page
    contains both prose and a table-like grid so the parser exercises
    its table-extraction path even on the fallback PyMuPDF backend.
    """
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.pdfgen import canvas
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "reportlab not installed; required for test PDF synthesis"
        ) from exc

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    width, height = LETTER
    text = c.beginText(72, height - 72)
    text.setFont("Helvetica-Bold", 14)
    text.textLine("Sample Hotel T-12 Operating Statement")
    text.setFont("Helvetica", 10)
    text.textLine("")
    text.textLine("Property: The Fondok Inn — Austin, TX")
    text.textLine("Period: Jan 2024 – Dec 2024")
    text.textLine("")
    text.textLine("Net Operating Income: $1,234,567")
    text.textLine("Occupancy: 74%")
    text.textLine("Average Daily Rate (ADR): $185.40")
    text.textLine("RevPAR: $137.20")
    text.textLine("")
    text.textLine("Department  | Revenue    | Expense   | Profit")
    text.textLine("Rooms       | 6,500,000  | 1,500,000 | 5,000,000")
    text.textLine("F&B         | 1,800,000  | 1,200,000 |   600,000")
    text.textLine("Other       |   200,000  |    80,000 |   120,000")
    c.drawText(text)
    c.showPage()
    c.save()
    return buf.getvalue()


@pytest.fixture(scope="session")
def sample_pdf_bytes() -> bytes:
    pdf_path = Path(__file__).parent / "fixtures" / "sample_t12.pdf"
    if pdf_path.exists():
        return pdf_path.read_bytes()
    body = _build_sample_pdf()
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    pdf_path.write_bytes(body)
    return body


@pytest.fixture
async def deal_id() -> str:
    """Insert a stub deal row directly so document FKs resolve."""
    from sqlalchemy import text

    from app.config import get_settings
    from app.database import dispose_engine, get_session_factory
    from app.migrations import run_startup_migrations

    # Make sure schema exists.
    await run_startup_migrations()

    settings = get_settings()
    factory = get_session_factory()
    from uuid import uuid4

    new_id = uuid4()
    async with factory() as session:
        await session.execute(
            text(
                """
                INSERT INTO deals (id, tenant_id, name, status, created_at, updated_at)
                VALUES (:id, :tenant, :name, 'Draft', :ts, :ts)
                """
            ),
            {
                "id": str(new_id),
                "tenant": settings.DEFAULT_TENANT_ID,
                "name": "Test Hotel",
                "ts": "2026-04-27 00:00:00",
            },
        )
        await session.commit()
    yield str(new_id)
    await dispose_engine()


# ─────────────────────────── parser ───────────────────────────


@pytest.mark.asyncio
async def test_parse_pymupdf_fallback(sample_pdf_bytes: bytes) -> None:
    """Without LLAMA_CLOUD_API_KEY the parser must use PyMuPDF."""
    # Defensive: ensure no leftover key in the env for this test.
    os.environ.pop("LLAMA_CLOUD_API_KEY", None)

    from app.extraction import parse_pdf

    parsed = await parse_pdf(sample_pdf_bytes, "sample_t12.pdf")
    assert parsed.filename == "sample_t12.pdf"
    assert parsed.total_pages >= 1
    assert parsed.pages, "no pages parsed"
    page_text = parsed.pages[0].text.lower()
    assert "net operating income" in page_text
    assert parsed.parser.startswith("pymupdf")
    assert parsed.content_hash and len(parsed.content_hash) == 64


@pytest.mark.asyncio
async def test_parse_xls_str_trend() -> None:
    """STR CoStar Trend reports ship as legacy .xls — each sheet
    becomes one ParsedPage, table grids preserved, parser='xlrd'.
    """
    from app.extraction import parse_document

    xls_path = Path(__file__).parent / "fixtures" / "sample_str_trend.xls"
    if not xls_path.exists():
        pytest.skip("sample_str_trend.xls fixture not present")

    body = xls_path.read_bytes()
    parsed = await parse_document(body, "sample_str_trend.xls")

    assert parsed.parser == "xlrd"
    assert parsed.total_pages >= 8, "expected multi-sheet workbook"
    assert all(p.metadata.get("sheet_name") for p in parsed.pages), (
        "every page must carry its sheet name"
    )
    # By Measure sheet (sheet 2) carries the headline Occ/ADR/RevPAR table
    by_measure = parsed.pages[1]
    assert "Occupancy" in by_measure.text or "occupancy" in by_measure.text.lower()
    assert by_measure.tables, "table grid must be preserved"
    assert parsed.content_hash and len(parsed.content_hash) == 64


@pytest.mark.asyncio
async def test_parse_unsupported_extension_raises() -> None:
    """Any extension other than the registered set is a hard ParseError —
    callers (the upload background task) catch this and mark the row
    PARSE_FAILED rather than silently dropping the file.
    """
    from app.extraction import ParseError, parse_document

    with pytest.raises(ParseError, match="unsupported file extension"):
        await parse_document(b"some-bytes", "notes.txt")


@pytest.mark.asyncio
async def test_parse_docx_extracts_paragraphs_and_tables() -> None:
    """Sam QA 2026-06-29: HMA Summary + Business Plan docs were
    failing as 'unsupported file extension .docx'. python-docx now
    handles them — paragraphs in document order, tables both as
    structured grids AND folded into the text stream so the LLM
    extractor sees them positionally.
    """
    pytest.importorskip("docx")
    from io import BytesIO

    from docx import Document  # type: ignore[import-untyped]

    from app.extraction import parse_document

    d = Document()
    d.add_heading("Anglers HMA Summary", 0)
    d.add_paragraph("Term: 25 years from 2018.")
    t = d.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Base fee"
    t.rows[0].cells[1].text = "3.0% of gross"
    t.rows[1].cells[0].text = "Incentive"
    t.rows[1].cells[1].text = "15% of GOP"
    buf = BytesIO()
    d.save(buf)
    body = buf.getvalue()

    parsed = await parse_document(body, "hma.docx")

    assert parsed.parser == "python-docx"
    assert parsed.total_pages == 1
    page = parsed.pages[0]
    assert "Anglers HMA Summary" in page.text
    assert "25 years" in page.text
    # Table folded into text + preserved as structured grid.
    assert "Base fee" in page.text
    assert "3.0% of gross" in page.text
    assert len(page.tables) == 1
    table = page.tables[0]
    assert ["Base fee", "3.0% of gross"] in table
    assert ["Incentive", "15% of GOP"] in table
    assert parsed.content_hash and len(parsed.content_hash) == 64


@pytest.mark.asyncio
async def test_parse_docx_corrupt_file_surfaces_clear_error() -> None:
    """A non-OOXML payload uploaded as .docx must surface a clean
    ParseError instead of crashing or hanging. (Sam QA: a couple of
    staged .docx files were corrupted 'not a zip file' — the no-
    silent-failures pass requires this to read as actionable.)
    """
    pytest.importorskip("docx")
    from app.extraction import ParseError, parse_document

    with pytest.raises(ParseError, match="python-docx failed to open"):
        await parse_document(b"not a real docx", "broken.docx")


def test_unknown_doc_type_falls_back_to_filename_hint() -> None:
    """Regression: when the router returns ``UNKNOWN`` (LLM rate-limit,
    credit-balance exhausted, off-list emission), the extraction
    pipeline must fall back to the filename hint instead of crashing
    on Pydantic enum validation.

    The fix lives at apps/worker/app/api/documents.py
    `_run_graph_extraction` — it now intersects the router output with
    the valid DocType enum members and falls back to ``_guess_doc_type``
    when there's no match. This test exercises both halves: the enum
    set and the filename hint.
    """
    from fondok_schemas import DocType

    from app.api.documents import _guess_doc_type

    valid = {dt.value for dt in DocType}
    assert "UNKNOWN" not in valid, (
        "UNKNOWN must NOT appear as a real DocType — it's a sentinel "
        "the router emits when it can't classify"
    )

    # The filename hint must always resolve to a valid DocType so the
    # fallback never re-introduces the original crash.
    for filename in [
        "sample_OM.pdf",
        "sample_T12.pdf",
        "sample_str_trend.xls",
        "sample_cbre_horizons.pdf",
        "sample_pnl_benchmark.pdf",
        "Random_File.pdf",  # falls back to T12 default
    ]:
        hint = _guess_doc_type(filename)
        assert hint in valid, (
            f"_guess_doc_type returned {hint!r} for {filename!r}, "
            f"which is not a valid DocType — the UNKNOWN fallback "
            f"would still crash"
        )


# ─────────────────────────── storage ───────────────────────────


@pytest.mark.asyncio
async def test_local_raw_store(tmp_path: Path) -> None:
    """LocalRawStore put/get roundtrip with a deterministic key."""
    from app.storage import LocalRawStore

    store = LocalRawStore(tmp_path / "raw")
    body = b"hello-world"
    key = await store.put(
        tenant_id="t1",
        deal_id="d1",
        content_hash="abc123",
        filename="thing.pdf",
        bytes_=body,
    )
    assert key.startswith("file://")
    assert await store.exists(key)
    got = await store.get(key)
    assert got == body


# ─────────────────────────── upload ───────────────────────────


@pytest.mark.asyncio
async def test_upload_pdf(sample_pdf_bytes: bytes, deal_id: str) -> None:
    """Upload writes a documents row, parks the file in the raw store,
    and returns immediately with the row at status ``PARSING``.

    The actual PDF parse + extraction runs as a background task so
    dense uploads don't blow through the proxy's HTTP timeout. Per-row
    fields populated synchronously: id, filename, content_hash,
    storage_key, status. Fields populated asynchronously after parse:
    page_count, parser, extraction_data.
    """
    from httpx import ASGITransport, AsyncClient

    from app.main import app
    from app.storage import get_raw_store, reset_raw_store_cache

    reset_raw_store_cache()  # pick up the test DOCUMENT_STORAGE_ROOT

    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test"
    ) as client:
        files = {
            "files": ("sample_t12.pdf", sample_pdf_bytes, "application/pdf"),
        }
        r = await client.post(f"/deals/{deal_id}/documents/upload", files=files)

    assert r.status_code == 201, r.text
    body = r.json()
    assert isinstance(body, list) and len(body) == 1
    rec = body[0]
    assert rec["filename"] == "sample_t12.pdf"
    # Synchronous parts of the upload:
    assert rec["status"] == "PARSING"
    assert rec["content_hash"] and len(rec["content_hash"]) == 64
    assert rec["storage_key"].startswith("file://")
    # Async parts haven't filled in yet (parse is in flight).
    assert rec["page_count"] is None
    assert rec["parser"] is None

    # File exists on disk where we said it would.
    store = get_raw_store()
    assert await store.exists(rec["storage_key"])


# ─────────────────────────── end-to-end ───────────────────────────


@pytest.mark.asyncio
async def test_extraction_flow_end_to_end(
    sample_pdf_bytes: bytes, deal_id: str
) -> None:
    """Upload → extract → poll → fetch extraction. Uses EVALS_MOCK."""
    from httpx import ASGITransport, AsyncClient

    from app.main import app
    from app.storage import reset_raw_store_cache

    reset_raw_store_cache()
    os.environ["EVALS_MOCK"] = "true"

    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test"
    ) as client:
        # Upload
        r = await client.post(
            f"/deals/{deal_id}/documents/upload",
            files={
                "files": ("sample_t12.pdf", sample_pdf_bytes, "application/pdf"),
            },
        )
        assert r.status_code == 201, r.text
        doc_id = r.json()[0]["id"]
        UUID(doc_id)  # validates shape

        # List
        r = await client.get(f"/deals/{deal_id}/documents")
        assert r.status_code == 200
        assert any(d["id"] == doc_id for d in r.json())

        # Upload now auto-chains parse → extract on its own background
        # task — no explicit POST /extract needed. We just poll the
        # extraction endpoint until status flips to EXTRACTED. (The
        # explicit /extract route is still available as a manual
        # re-run mechanism but doesn't need to fire on every upload.)
        final_status = None
        for _ in range(80):  # ~8s budget — async parse + extract
            r = await client.get(
                f"/deals/{deal_id}/documents/{doc_id}/extraction"
            )
            assert r.status_code == 200, r.text
            final_status = r.json()["status"]
            if final_status in ("EXTRACTED", "FAILED", "PARSE_FAILED"):
                break
            await asyncio.sleep(0.1)

        assert final_status == "EXTRACTED", (
            f"expected EXTRACTED, got {final_status}: {r.json()}"
        )

        body = r.json()
        assert body["fields"], "mock extraction should populate fields"
        # Mock payload always emits noi_year_1.
        assert any(f["field_name"] == "noi_year_1" for f in body["fields"])
        # The mock payload starts at overall 0.9, but the citation
        # verifier runs after extraction and promotes any field it
        # confirms verbatim against the source page (Sam QA 2026-05-14
        # critic-promote). The mock T-12 fixture carries "Net Operating
        # Income: $1,234,567" so noi_year_1 verifies MATCH and is
        # floored to 0.98 — the rolled-up overall lands at-or-above the
        # raw 0.9. Assert the floor, not an exact value.
        assert body["confidence_report"]["overall"] >= 0.9


# ─────────────────────────── graph ───────────────────────────


def test_graph_compiles_with_real_nodes() -> None:
    """The real graph (with bound node functions) must compile.

    Distinct from ``test_smoke.test_graph_compiles`` because we assert
    on the node set — the wiring must include every pipeline stage.
    """
    from app.graph import build_graph

    g = build_graph()
    assert g is not None

    # The compiled graph exposes its node set on .nodes.
    expected = {
        "route",
        "extract",
        "normalize",
        "gate1_review",
        "run_engines",
        "analyze",
        "variance",
        "gate2_review",
        "finalize",
    }
    assert expected.issubset(set(g.nodes))


# ─────────────────── content-hash extraction cache ──────────────────
#
# Sam cost-opt 2026-07: every upload used to run Router → Extractor →
# Normalizer → Verifier against Claude even when the exact same bytes
# were extracted last week. At $6-10 per deal in LLM tokens this is
# the biggest waste. The cache short-circuits identical-content
# re-uploads to zero LLM cost. Safety gates: same-tenant only, same
# pipeline version only, source row must be status=EXTRACTED. These
# tests seed the DB directly so they exercise the cache-check gate
# without going through the upload endpoint.


async def _seed_deal_and_document(
    *,
    tenant_id: str,
    content_hash: str,
    filename: str = "cached.pdf",
    status: str = "EXTRACTED",
) -> tuple[str, str]:
    """Insert a stub deal + a stub document with the given content_hash.

    Returns ``(deal_id, doc_id)``. The document row carries a minimal
    ``extraction_data`` payload (one page with realistic text) so the
    downstream USALI-scoring / structural-recognizer post-processing
    inside ``_run_extraction_pipeline_inner`` has something to chew on
    without crashing.
    """
    import json as _json
    from uuid import uuid4

    from sqlalchemy import text

    from app.database import get_session_factory
    from app.migrations import run_startup_migrations

    await run_startup_migrations()
    factory = get_session_factory()

    deal_id = str(uuid4())
    doc_id = str(uuid4())
    extraction_data = {
        "parser": "test-fixture",
        "total_pages": 1,
        "content_hash": content_hash,
        "parsed_at": "2026-07-01T00:00:00+00:00",
        "pages": [
            {"page_num": 1, "text": "Test doc", "tables": [], "metadata": {}}
        ],
    }
    async with factory() as session:
        await session.execute(
            text(
                """
                INSERT INTO deals (id, tenant_id, name, status,
                                   created_at, updated_at)
                VALUES (:id, :tenant, :name, 'Draft', :ts, :ts)
                """
            ),
            {
                "id": deal_id,
                "tenant": tenant_id,
                "name": "Cache Test Hotel",
                "ts": "2026-07-01 00:00:00",
            },
        )
        await session.execute(
            text(
                """
                INSERT INTO documents (
                    id, deal_id, tenant_id, filename, doc_type, status,
                    uploaded_at, content_hash, storage_key, size_bytes,
                    page_count, parser, extraction_data
                ) VALUES (
                    :id, :deal, :tenant, :filename, :dt, :status,
                    :uploaded_at, :h, :sk, :size,
                    :pages, :parser, :ed
                )
                """
            ),
            {
                "id": doc_id,
                "deal": deal_id,
                "tenant": tenant_id,
                "filename": filename,
                "dt": "T12",
                "status": status,
                "uploaded_at": "2026-07-01 00:00:00",
                "h": content_hash,
                "sk": "file:///tmp/dummy",
                "size": 1024,
                "pages": 1,
                "parser": "test-fixture",
                "ed": _json.dumps(extraction_data),
            },
        )
        await session.commit()
    return deal_id, doc_id


async def _seed_extraction_result(
    *,
    doc_id: str,
    deal_id: str,
    tenant_id: str,
    agent_version: str,
) -> str:
    """Insert a canned extraction_results row for cache-lookup tests."""
    import json as _json
    from uuid import uuid4

    from sqlalchemy import text

    from app.database import get_session_factory

    factory = get_session_factory()
    ext_id = str(uuid4())
    fields = [
        {
            "field_name": "noi_year_1",
            "value": 999999.0,
            "unit": "USD",
            "source_page": 1,
            "confidence": 0.9,
            "raw_text": "Cached NOI",
        }
    ]
    confidence = {
        "overall": 0.9,
        "by_field": {"noi_year_1": 0.9},
        "low_confidence_fields": [],
        "requires_human_review": False,
    }
    async with factory() as session:
        await session.execute(
            text(
                """
                INSERT INTO extraction_results (
                    id, document_id, deal_id, tenant_id,
                    fields, confidence_report, agent_version, created_at
                ) VALUES (
                    :id, :doc, :deal, :tenant,
                    :fields, :cr, :ver, :created
                )
                """
            ),
            {
                "id": ext_id,
                "doc": doc_id,
                "deal": deal_id,
                "tenant": tenant_id,
                "fields": _json.dumps(fields),
                "cr": _json.dumps(confidence),
                "ver": agent_version,
                "created": "2026-07-01 00:00:00",
            },
        )
        await session.commit()
    return ext_id


@pytest.mark.asyncio
async def test_extraction_cache_hit_reuses_prior_result_no_llm_call() -> None:
    """Same tenant, same content_hash, same pipeline version →
    ``_run_extraction_pipeline`` skips ``_run_graph_extraction`` entirely
    and clones the prior fields + confidence_report into a new row.

    Guardrail: we patch ``_run_graph_extraction`` to raise so any code
    path that actually invokes the LLM chain would blow the test up.
    """
    from unittest.mock import AsyncMock, patch

    from sqlalchemy import text

    from app.api import documents as docs_module
    from app.config import get_settings
    from app.database import get_session_factory

    # EVALS_MOCK would ALSO short-circuit — clear it so we prove the
    # cache branch itself is doing the work.
    os.environ.pop("EVALS_MOCK", None)
    docs_module._reset_extraction_cache_metrics()

    settings = get_settings()
    tenant_id = settings.DEFAULT_TENANT_ID
    content_hash = "a" * 64
    src_agent_version = f"router:T12;extractor;pv={docs_module.EXTRACTION_PIPELINE_VERSION}"

    # Seed the source (already-extracted) doc + its result row.
    src_deal, src_doc = await _seed_deal_and_document(
        tenant_id=tenant_id,
        content_hash=content_hash,
        filename="source.pdf",
    )
    src_ext_id = await _seed_extraction_result(
        doc_id=src_doc,
        deal_id=src_deal,
        tenant_id=tenant_id,
        agent_version=src_agent_version,
    )

    # Seed the NEW doc (same content_hash, freshly UPLOADED — awaiting extraction).
    new_deal, new_doc = await _seed_deal_and_document(
        tenant_id=tenant_id,
        content_hash=content_hash,
        filename="dupe.pdf",
        status="UPLOADED",
    )

    # Guardrail: if the cache-check gate misses and we fall through to
    # the real extractor, this mock explodes and the test fails loudly.
    boom = AsyncMock(
        side_effect=AssertionError(
            "cache HIT should have short-circuited — _run_graph_extraction "
            "must NOT be called"
        )
    )
    with patch.object(docs_module, "_run_graph_extraction", boom):
        await docs_module._run_extraction_pipeline(
            deal_id=new_deal, doc_id=new_doc, tenant_id=tenant_id,
        )

    boom.assert_not_awaited()

    # The new doc landed EXTRACTED with a cloned extraction_results row.
    factory = get_session_factory()
    async with factory() as session:
        r = (
            await session.execute(
                text("SELECT status FROM documents WHERE id = :id"),
                {"id": new_doc},
            )
        ).first()
        assert r._mapping["status"] == "EXTRACTED"

        r = (
            await session.execute(
                text(
                    "SELECT id, agent_version, fields FROM extraction_results "
                    "WHERE document_id = :id"
                ),
                {"id": new_doc},
            )
        ).first()
        assert r is not None, "cache HIT must still persist a per-doc row"
        # Distinct row (tenant-scoped per-doc queries still work).
        assert str(r._mapping["id"]) != src_ext_id
        assert r._mapping["agent_version"] == src_agent_version
        # Cloned field payload — same value the source row carried.
        cloned = r._mapping["fields"]
        if isinstance(cloned, str):
            import json as _json

            cloned = _json.loads(cloned)
        assert cloned[0]["field_name"] == "noi_year_1"
        assert cloned[0]["value"] == 999999.0

    metrics = docs_module.get_extraction_cache_metrics()
    assert metrics["per_tenant"][tenant_id]["hits"] == 1
    assert metrics["per_tenant"][tenant_id]["misses"] == 0


@pytest.mark.asyncio
async def test_extraction_cache_miss_on_different_content_hash() -> None:
    """Different content_hash → no cache hit → extractor runs."""
    from unittest.mock import AsyncMock, patch

    from app.api import documents as docs_module
    from app.config import get_settings

    os.environ.pop("EVALS_MOCK", None)
    docs_module._reset_extraction_cache_metrics()

    settings = get_settings()
    tenant_id = settings.DEFAULT_TENANT_ID

    # Seed a prior extracted doc with hash A.
    src_deal, src_doc = await _seed_deal_and_document(
        tenant_id=tenant_id, content_hash="a" * 64, filename="A.pdf",
    )
    await _seed_extraction_result(
        doc_id=src_doc,
        deal_id=src_deal,
        tenant_id=tenant_id,
        agent_version=f"router:T12;extractor;pv={docs_module.EXTRACTION_PIPELINE_VERSION}",
    )

    # New doc with hash B — must MISS.
    new_deal, new_doc = await _seed_deal_and_document(
        tenant_id=tenant_id,
        content_hash="b" * 64,
        filename="B.pdf",
        status="UPLOADED",
    )

    async def _stub_extractor(**_kwargs):
        # Return a minimal envelope that mimics the real extractor's
        # shape so the post-processing tail (structural override, USALI
        # scoring) can run cleanly.
        return (
            [
                {
                    "field_name": "noi_year_1",
                    "value": 111.0,
                    "unit": "USD",
                    "source_page": 1,
                    "confidence": 0.9,
                    "raw_text": "fresh",
                }
            ],
            {
                "overall": 0.9,
                "by_field": {"noi_year_1": 0.9},
                "low_confidence_fields": [],
                "requires_human_review": False,
            },
            "router:T12;extractor",
            "T12",
        )

    mock_extractor = AsyncMock(side_effect=_stub_extractor)
    with patch.object(docs_module, "_run_graph_extraction", mock_extractor):
        await docs_module._run_extraction_pipeline(
            deal_id=new_deal, doc_id=new_doc, tenant_id=tenant_id,
        )

    mock_extractor.assert_awaited_once()

    metrics = docs_module.get_extraction_cache_metrics()
    assert metrics["per_tenant"][tenant_id]["hits"] == 0
    assert metrics["per_tenant"][tenant_id]["misses"] == 1


@pytest.mark.asyncio
async def test_extraction_cache_miss_on_different_agent_version() -> None:
    """Same content_hash but the prior row was written by an OLDER pipeline
    version (``pv=v0``) → MISS. Prevents stale-cache serves after an
    extractor redeploy.
    """
    from unittest.mock import AsyncMock, patch

    from app.api import documents as docs_module
    from app.config import get_settings

    os.environ.pop("EVALS_MOCK", None)
    docs_module._reset_extraction_cache_metrics()

    settings = get_settings()
    tenant_id = settings.DEFAULT_TENANT_ID
    content_hash = "c" * 64

    src_deal, src_doc = await _seed_deal_and_document(
        tenant_id=tenant_id, content_hash=content_hash, filename="old.pdf",
    )
    # Legacy pipeline version — must NOT satisfy the lookup.
    await _seed_extraction_result(
        doc_id=src_doc,
        deal_id=src_deal,
        tenant_id=tenant_id,
        agent_version="router:T12;extractor;pv=v0",
    )

    new_deal, new_doc = await _seed_deal_and_document(
        tenant_id=tenant_id,
        content_hash=content_hash,
        filename="new.pdf",
        status="UPLOADED",
    )

    async def _stub_extractor(**_kwargs):
        return (
            [
                {
                    "field_name": "noi_year_1",
                    "value": 222.0,
                    "unit": "USD",
                    "source_page": 1,
                    "confidence": 0.9,
                    "raw_text": "fresh",
                }
            ],
            {
                "overall": 0.9,
                "by_field": {"noi_year_1": 0.9},
                "low_confidence_fields": [],
                "requires_human_review": False,
            },
            "router:T12;extractor",
            "T12",
        )

    mock_extractor = AsyncMock(side_effect=_stub_extractor)
    with patch.object(docs_module, "_run_graph_extraction", mock_extractor):
        await docs_module._run_extraction_pipeline(
            deal_id=new_deal, doc_id=new_doc, tenant_id=tenant_id,
        )

    mock_extractor.assert_awaited_once()

    metrics = docs_module.get_extraction_cache_metrics()
    assert metrics["per_tenant"][tenant_id]["misses"] == 1
    assert metrics["per_tenant"][tenant_id]["hits"] == 0


@pytest.mark.asyncio
async def test_extraction_cache_isolates_tenants() -> None:
    """A cached extraction on tenant A must NEVER short-circuit an
    upload on tenant B — hard security requirement.
    """
    from unittest.mock import AsyncMock, patch

    from app.api import documents as docs_module

    os.environ.pop("EVALS_MOCK", None)
    docs_module._reset_extraction_cache_metrics()

    tenant_a = "11111111-1111-1111-1111-111111111111"
    tenant_b = "22222222-2222-2222-2222-222222222222"
    content_hash = "d" * 64
    src_agent_version = (
        f"router:T12;extractor;pv={docs_module.EXTRACTION_PIPELINE_VERSION}"
    )

    # Tenant A: seed an already-extracted doc with the shared content_hash.
    a_deal, a_doc = await _seed_deal_and_document(
        tenant_id=tenant_a, content_hash=content_hash, filename="A.pdf",
    )
    await _seed_extraction_result(
        doc_id=a_doc,
        deal_id=a_deal,
        tenant_id=tenant_a,
        agent_version=src_agent_version,
    )

    # Tenant B: new upload of the SAME bytes. Must NOT cache-hit.
    b_deal, b_doc = await _seed_deal_and_document(
        tenant_id=tenant_b,
        content_hash=content_hash,
        filename="B.pdf",
        status="UPLOADED",
    )

    async def _stub_extractor(**_kwargs):
        return (
            [
                {
                    "field_name": "noi_year_1",
                    "value": 333.0,
                    "unit": "USD",
                    "source_page": 1,
                    "confidence": 0.9,
                    "raw_text": "fresh-b",
                }
            ],
            {
                "overall": 0.9,
                "by_field": {"noi_year_1": 0.9},
                "low_confidence_fields": [],
                "requires_human_review": False,
            },
            "router:T12;extractor",
            "T12",
        )

    mock_extractor = AsyncMock(side_effect=_stub_extractor)
    with patch.object(docs_module, "_run_graph_extraction", mock_extractor):
        await docs_module._run_extraction_pipeline(
            deal_id=b_deal, doc_id=b_doc, tenant_id=tenant_b,
        )

    # The extractor ran for tenant B (no cross-tenant cache hit).
    mock_extractor.assert_awaited_once()

    metrics = docs_module.get_extraction_cache_metrics()
    # Tenant B took a miss; tenant A never appears in the counters (its
    # doc was seeded directly, not routed through the pipeline).
    assert metrics["per_tenant"][tenant_b]["misses"] == 1
    assert metrics["per_tenant"].get(tenant_b, {}).get("hits", 0) == 0
    assert tenant_a not in metrics["per_tenant"]


@pytest.mark.asyncio
async def test_extraction_cache_disabled_flag_bypasses_lookup() -> None:
    """``EXTRACTION_CACHE_ENABLED=False`` disables the cache entirely —
    the extractor runs even when a valid cached row exists. Ops uses
    this as a kill switch when debugging a suspected stale-cache issue.
    """
    from unittest.mock import AsyncMock, patch

    from app.api import documents as docs_module
    from app.config import Settings, get_settings

    os.environ.pop("EVALS_MOCK", None)
    docs_module._reset_extraction_cache_metrics()

    settings = get_settings()
    tenant_id = settings.DEFAULT_TENANT_ID
    content_hash = "e" * 64
    src_agent_version = (
        f"router:T12;extractor;pv={docs_module.EXTRACTION_PIPELINE_VERSION}"
    )

    src_deal, src_doc = await _seed_deal_and_document(
        tenant_id=tenant_id, content_hash=content_hash, filename="src.pdf",
    )
    await _seed_extraction_result(
        doc_id=src_doc,
        deal_id=src_deal,
        tenant_id=tenant_id,
        agent_version=src_agent_version,
    )

    new_deal, new_doc = await _seed_deal_and_document(
        tenant_id=tenant_id,
        content_hash=content_hash,
        filename="new.pdf",
        status="UPLOADED",
    )

    # Flip the flag OFF via a patched settings instance.
    disabled = Settings(EXTRACTION_CACHE_ENABLED=False)

    async def _stub_extractor(**_kwargs):
        return (
            [
                {
                    "field_name": "noi_year_1",
                    "value": 444.0,
                    "unit": "USD",
                    "source_page": 1,
                    "confidence": 0.9,
                    "raw_text": "fresh",
                }
            ],
            {
                "overall": 0.9,
                "by_field": {"noi_year_1": 0.9},
                "low_confidence_fields": [],
                "requires_human_review": False,
            },
            "router:T12;extractor",
            "T12",
        )

    mock_extractor = AsyncMock(side_effect=_stub_extractor)
    with patch.object(docs_module, "get_settings", return_value=disabled), \
         patch.object(docs_module, "_run_graph_extraction", mock_extractor):
        await docs_module._run_extraction_pipeline(
            deal_id=new_deal, doc_id=new_doc, tenant_id=tenant_id,
        )

    # Cache was reachable but the flag disabled it — extractor ran.
    mock_extractor.assert_awaited_once()

    metrics = docs_module.get_extraction_cache_metrics()
    assert metrics["per_tenant"][tenant_id]["misses"] == 1
    assert metrics["per_tenant"][tenant_id]["hits"] == 0
